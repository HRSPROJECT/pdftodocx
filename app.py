import streamlit as st
import io
from docx import Document
import fitz  # PyMuPDF
import base64
import pdfplumber
from PIL import Image
import tempfile

# CSS to hide Streamlit elements
hide_st_style = """
            <style>
            #MainMenu {visibility: hidden;}
            footer {visibility: hidden;}
            header {visibility: hidden;}
            </style>
            """

st.markdown(hide_st_style, unsafe_allow_html=True)


def create_theme_css(theme):
    """Creates the CSS for the dark or light theme."""
    if theme == "dark":
        return """
            <style>
                body {
                    background-color: #1a1a1a;
                    color: white;
                }
                .stButton>button {
                    background-color: #333;
                    color: white;
                    border: 1px solid #555;
                }
                .stButton>button:hover {
                    background-color: #555;
                }
                 .css-1q8dd3f {
                     background-color: #222 !important;
                    border-color: #555 !important;
                 }
                  .css-2y7684 {
                    background-color: #222 !important;
                    color: white;
                }
                .stTextInput > div > div > input{
                    background-color: #222;
                    color: white;
                }
                .stRadio > div > label {
                    color: white;
                }
               .css-1h5jfnv p{
                  color: white;
                }
                .css-10trblm > div > div > input{
                    background-color: #222;
                    color: white;
                }
                .css-10trblm > div > div > input::placeholder{
                   color: #666;
                }
                 .css-keje6a, .css-13sd3d5 {
                     background-color: #222 !important;
                 }
            </style>
            """
    else:
        return """
            <style>
                body {
                    background-color: #f0f0f0;
                    color: black;
                }
                .stButton>button {
                    background-color: #e0e0e0;
                    color: black;
                    border: 1px solid #ccc;
                 }
                .stButton>button:hover {
                    background-color: #ccc;
                }
                .css-1q8dd3f {
                     background-color: #fff !important;
                    border-color: #ccc !important;
                 }
                 .css-2y7684 {
                    background-color: #fff !important;
                    color: black;
                }
                .stTextInput > div > div > input{
                  background-color: white;
                  color: black;
               }
               .stRadio > div > label {
                  color: black;
               }
                .css-1h5jfnv p{
                  color: black;
                }
                .css-10trblm > div > div > input{
                    background-color: white;
                    color: black;
                }
                 .css-10trblm > div > div > input::placeholder{
                   color: #999;
                 }
                  .css-keje6a, .css-13sd3d5{
                      background-color: #fff !important;
                  }
            </style>
            """


# Function to create the navigation bar
def create_navbar(theme):
    navbar_style = f"""
        <nav style="background-color: {'#333' if theme == 'dark' else '#f0f0f0'}; padding: 10px; display: flex; justify-content: space-between; align-items: center;">
            <span style="font-size: 20px; font-weight: bold; color: {'white' if theme == 'dark' else 'black'};">PDF to DOCX Converter</span>
            <a href="https://hrsproject.github.io/home/" target="_blank" style="color: {'white' if theme == 'dark' else 'black'}; text-decoration: none; padding: 8px 12px; border: 1px solid {'white' if theme == 'dark' else 'black'}; border-radius: 4px;">Explore</a>
        </nav>
    """
    st.markdown(navbar_style, unsafe_allow_html=True)


# Function to create a download button that includes preview
def create_download_button(docx_stream, file_name, preview_text, theme):
    preview_html = f"""
            <div style="padding: 10px; margin-bottom: 10px; border: 1px solid {'#555' if theme == 'dark' else '#ccc'}; background-color: {'#222' if theme == 'dark' else '#fff'}; color: {'white' if theme == 'dark' else 'black'}; border-radius: 5px; max-height: 300px; overflow-y: auto;">
                <p style="font-size: 14px; white-space: pre-wrap;">{preview_text}</p>
            </div>
    """
    st.markdown(preview_html, unsafe_allow_html=True)
    st.download_button(
        label="Download DOCX",
        data=docx_stream,
        file_name=file_name,
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


# Function for OCR text extraction
def extract_text_from_pdf_ocr(pdf_file):
    """Extracts text from a PDF file using OCR."""
    text = ""
    try:
        with pdfplumber.open(pdf_file) as pdf:
            for page in pdf.pages:
                text += page.extract_text()
    except Exception as e:
        st.error(f"Error reading PDF: {e}")
        return None
    return text

# Function for direct PDF to DOCX conversion
def convert_pdf_to_docx(pdf_file):
    """Converts a PDF directly to DOCX without OCR (non-searchable text)."""
    try:
        pdf_document = fitz.open(stream=pdf_file.read(), filetype="pdf")
        document = Document()
        for page_num in range(len(pdf_document)):
            page = pdf_document[page_num]
            pix = page.get_pixmap() #Gets pixmap of each page
            image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples) #Makes image from pixmap
            temp_file = tempfile.NamedTemporaryFile(suffix='.png', delete=False) #Creates a temporary file for the image
            image.save(temp_file, format='png') #saves the temp file as png
            document.add_picture(temp_file.name, width=8) #Adds the file to the docx document
            temp_file.close() #closes the temp file
            temp_file.unlink(temp_file.name) #Delete the file
        return document
    except Exception as e:
        st.error(f"Error converting PDF directly to DOCX: {e}")
        return None

# Function to create a DOCX document
def create_docx(text):
    """Creates a DOCX document from the given text."""
    document = Document()
    document.add_paragraph(text)
    return document


def main():
    # Initialize theme in session state
    if "theme" not in st.session_state:
       st.session_state.theme = "light"

    # Theme toggle
    col1, col2 = st.columns(2)
    with col2:
        if st.button("Toggle Dark Mode" if st.session_state.theme == "light" else "Toggle Light Mode"):
            st.session_state.theme = "dark" if st.session_state.theme == "light" else "light"

    # Apply theme CSS
    st.markdown(create_theme_css(st.session_state.theme), unsafe_allow_html=True)
    create_navbar(st.session_state.theme)

    st.title("PDF to DOCX Converter")
    ocr_option = st.radio("Choose Conversion Option:", ["Perform OCR (Searchable Text)", "Direct Conversion (Non-searchable Text)"])
    uploaded_file = st.file_uploader("Upload a PDF file", type="pdf")


    if uploaded_file is not None:
        with st.spinner("Processing..."):
             if ocr_option == "Perform OCR (Searchable Text)":
                 pdf_text = extract_text_from_pdf_ocr(uploaded_file)
                 if pdf_text:
                     docx_document = create_docx(pdf_text)
                     docx_stream = io.BytesIO()
                     docx_document.save(docx_stream)
                     docx_stream.seek(0)

                     # Provide preview and download button
                     create_download_button(docx_stream, "converted.docx", pdf_text, st.session_state.theme)

             else: # Direct conversion without OCR
                docx_document = convert_pdf_to_docx(uploaded_file)
                if docx_document:
                    docx_stream = io.BytesIO()
                    docx_document.save(docx_stream)
                    docx_stream.seek(0)
                    create_download_button(docx_stream, "converted.docx", "No text preview for this conversion type.", st.session_state.theme)

    st.markdown("Created by [Your Name] with Streamlit")

if __name__ == "__main__":
    main()
